"""Ingest HSEN (Harmonized System Explanatory Notes) into the KG.

HSEN is WCO-published interpretive commentary on the HS nomenclature. It is
*the* most-cited classification authority short of the legal text itself.
HMRC and tribunals routinely cite HSEN when classification is contested.

Authority tier: 2 (binding-strength interpretive guidance). Not legal text
(that's T1: chapter/section notes + GIRs), but the highest persuasive
authority for classification. Above ATARs because ATARs are product-specific
whereas HSEN is the canonical reading of the nomenclature itself.

Per chapter doc we extract:
- GENERAL commentary -> one edge per chapter
- Per-heading explanatory notes (NN.NN - title) -> one edge per heading
- We SKIP the "Notes." and "Subheading Note." sections because those are
  already in our KG as decomposed Chapter/Section Note rules (tier 1).

Source string: 'hsen:ch64', 'hsen:section_xv', etc.

Cost: zero (no LLM involved - pure structural parsing).
"""
from __future__ import annotations

import os
import re
import sys
from pathlib import Path

from dotenv import load_dotenv

# Load env
for p in [Path(__file__).parent / ".env",
          Path(__file__).parent.parent / ".env",
          Path(os.environ["AI_FAN_OUT_ENV_FILE"]) if os.environ.get("AI_FAN_OUT_ENV_FILE") else None]:
    if p is not None and p.exists():
        load_dotenv(p)
        break

import psycopg
from psycopg.rows import dict_row
from docx import Document
try:
    from .evidence_labels import edge_labels
except ImportError:
    from evidence_labels import edge_labels

DSN = os.environ.get("TARIFF_DB_DSN", "postgresql:///tariff_db")
HSEN_DIR = Path(os.environ.get("HSEN_DIR", "data/hsen"))
HSEN_TIER = 3  # authoritative interpretive guidance, not binding legal text

# Match "64.01 - Waterproof footwear..." style heading-level titles
HEADING_RE = re.compile(r"^(\d{2})\.(\d{2})\s*[-–—]\s*(.+?)$")
# Match SECTION I, SECTION II, ... (roman numerals)
SECTION_RE = re.compile(r"^SECTION\s+([IVXLC]+)\b", re.IGNORECASE)


def _flat_chapter(code_two: str) -> str:
    return f"{code_two}00000000"


def with_actor(conn, actor: str):
    with conn.cursor() as cur:
        cur.execute("SELECT set_config('kg.actor', %s, false)", (actor,))
    conn.commit()


def _commodities_for_chapter(conn, ch2: str) -> list[str]:
    """All 10-digit live commodity codes under a chapter."""
    with conn.cursor() as cur:
        cur.execute(
            """
            SELECT goods_nomenclature_item_id AS code
            FROM uk.goods_nomenclatures
            WHERE producline_suffix = '80'
              AND validity_end_date IS NULL
              AND SUBSTRING(goods_nomenclature_item_id, 1, 2) = %s
            """,
            (ch2,),
        )
        return [r["code"] for r in cur.fetchall()]


def _commodities_for_heading(conn, heading4: str) -> list[str]:
    """All 10-digit live commodity codes under a heading (4-digit prefix)."""
    with conn.cursor() as cur:
        cur.execute(
            """
            SELECT goods_nomenclature_item_id AS code
            FROM uk.goods_nomenclatures
            WHERE producline_suffix = '80'
              AND validity_end_date IS NULL
              AND SUBSTRING(goods_nomenclature_item_id, 1, 4) = %s
            """,
            (heading4,),
        )
        return [r["code"] for r in cur.fetchall()]


def parse_chapter_doc(path: Path) -> tuple[str | None, str, dict[str, tuple[str, str]]]:
    """Returns (chapter_two_digit, general_body, {heading4: (title, body)}).

    chapter_two_digit is None for the sections doc.
    """
    doc = Document(path)
    paras = doc.paragraphs
    ch2 = None
    general_body: list[str] = []
    headings: dict[str, tuple[str, list[str]]] = {}  # heading4 -> (title, body_paras)

    state = "preamble"  # preamble | notes | subheading_notes | general | heading
    current_heading: str | None = None

    for p in paras:
        text = p.text.strip()
        style = p.style.name if p.style else ""

        # Detect chapter number from "Chapter NN" or filename
        if state == "preamble" and style.startswith("Heading") and text.lower().startswith("chapter"):
            m = re.match(r"chapter\s+(\d{1,2})", text, re.IGNORECASE)
            if m:
                ch2 = m.group(1).zfill(2)
            continue

        if not text:
            continue

        # Section transitions
        low = text.lower()
        if low.startswith("notes.") or low == "notes":
            state = "notes"
            continue
        if low.startswith("subheading note") or low.startswith("sub-heading note"):
            state = "subheading_notes"
            continue
        if text == "GENERAL":
            state = "general"
            continue

        # Per-heading markers
        m = HEADING_RE.match(text)
        if m and style.startswith("Heading"):
            ch_part = m.group(1)
            sub = m.group(2)
            current_heading = f"{ch_part}{sub}"
            title = m.group(3).strip().rstrip(".")
            headings[current_heading] = (title, [])
            state = "heading"
            continue

        # Accumulate body text
        if state == "general":
            general_body.append(text)
        elif state == "heading" and current_heading is not None:
            headings[current_heading][1].append(text)
        # Skip notes/subheading_notes - already in KG as decomposed T1 rules

    flat_headings = {h: (title, "\n".join(body)) for h, (title, body) in headings.items() if body}
    return ch2, "\n".join(general_body), flat_headings


def parse_sections_doc(path: Path) -> dict[str, tuple[str, str]]:
    """Sections doc structure: SECTION I, II, ... each has GENERAL commentary.

    Returns {roman_numeral: (heading, body)}.
    """
    doc = Document(path)
    sections: dict[str, list[str]] = {}
    current: str | None = None
    section_titles: dict[str, str] = {}
    capture = False
    next_is_title = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = SECTION_RE.match(text)
        if m:
            current = m.group(1).upper()
            sections.setdefault(current, [])
            capture = True
            next_is_title = True
            continue
        if next_is_title and current:
            section_titles[current] = text[:200]
            next_is_title = False
            continue
        if capture and current:
            sections[current].append(text)
    return {s: (section_titles.get(s, f"Section {s}"), "\n".join(b)) for s, b in sections.items() if b}


# Roman numeral -> ordinal map (HS has 21 sections)
ROMAN_TO_INT = {
    "I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8,
    "IX": 9, "X": 10, "XI": 11, "XII": 12, "XIII": 13, "XIV": 14, "XV": 15,
    "XVI": 16, "XVII": 17, "XVIII": 18, "XIX": 19, "XX": 20, "XXI": 21,
}

# HS Section -> chapter range. Source: WCO HS structure.
SECTION_CHAPTERS = {
    1: range(1, 6), 2: range(6, 15), 3: [15], 4: range(16, 25), 5: range(25, 28),
    6: range(28, 39), 7: range(39, 41), 8: range(41, 44), 9: range(44, 47),
    10: range(47, 50), 11: range(50, 64), 12: range(64, 68), 13: range(68, 71),
    14: [71], 15: range(72, 84), 16: range(84, 86), 17: range(86, 90),
    18: range(90, 93), 19: [93], 20: range(94, 97), 21: [97],
}


def _commodities_for_section(conn, roman: str) -> list[str]:
    n = ROMAN_TO_INT.get(roman.upper())
    if not n:
        return []
    chapters = SECTION_CHAPTERS.get(n, [])
    if not chapters:
        return []
    with conn.cursor() as cur:
        codes = []
        for ch in chapters:
            cur.execute(
                """
                SELECT goods_nomenclature_item_id AS code
                FROM uk.goods_nomenclatures
                WHERE producline_suffix = '80'
                  AND validity_end_date IS NULL
                  AND SUBSTRING(goods_nomenclature_item_id, 1, 2) = %s
                """,
                (f"{ch:02d}",),
            )
            codes.extend(r["code"] for r in cur.fetchall())
        return codes


def upsert_edge(conn, edge_id: str, etype: str, scope: str, title: str, body: str, source: str, tier: int, provenance: dict) -> None:
    import json
    use_scopes, evidence_roles = edge_labels(etype, tier, source=source, edge_id=edge_id, scope=scope)
    with conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO kg.kg_edges (id, type, scope, title, body, source, authority_tier, use_scopes, evidence_roles, provenance)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s::text[], %s::text[], %s::jsonb)
            ON CONFLICT (id) DO UPDATE
            SET title = EXCLUDED.title,
                body = EXCLUDED.body,
                source = EXCLUDED.source,
                authority_tier = EXCLUDED.authority_tier,
                use_scopes = EXCLUDED.use_scopes,
                evidence_roles = EXCLUDED.evidence_roles,
                provenance = EXCLUDED.provenance,
                updated_at = now()
            """,
            (edge_id, etype, scope, title[:300], body[:30000], source, tier, use_scopes, evidence_roles, json.dumps(provenance)),
        )


def link_edge_to_commodities(conn, edge_id: str, codes: list[str]) -> int:
    if not codes:
        return 0
    with conn.cursor() as cur:
        cur.executemany(
            """
            INSERT INTO kg.kg_edge_commodities (edge_id, commodity_code)
            VALUES (%s, %s)
            ON CONFLICT DO NOTHING
            """,
            [(edge_id, c) for c in codes],
        )
    return len(codes)


def ingest_chapter(conn, path: Path) -> tuple[int, int]:
    """Returns (edges_upserted, links_inserted)."""
    ch2, general, headings = parse_chapter_doc(path)
    if not ch2:
        print(f"  could not detect chapter in {path.name}")
        return 0, 0
    edges = 0
    links = 0
    # GENERAL commentary -> one chapter-scoped edge
    if general.strip():
        edge_id = f"hsen:ch{ch2}:general"
        upsert_edge(
            conn, edge_id, "hsen_general", f"chapter:{ch2}",
            f"HSEN Chapter {ch2} - General",
            general,
            f"hsen:ch{ch2}",
            HSEN_TIER,
            {"source_type": "hsen", "chapter": ch2, "section": "GENERAL",
             "doc": path.name, "doc_part": "GENERAL"},
        )
        edges += 1
        codes = _commodities_for_chapter(conn, ch2)
        links += link_edge_to_commodities(conn, edge_id, codes)
    # Per-heading commentary
    for heading4, (title, body) in headings.items():
        edge_id = f"hsen:h{heading4}"
        upsert_edge(
            conn, edge_id, "hsen_heading", f"heading:{heading4}",
            f"HSEN {heading4[:2]}.{heading4[2:]} - {title}",
            body,
            f"hsen:ch{ch2}",
            HSEN_TIER,
            {"source_type": "hsen", "chapter": ch2, "heading": heading4,
             "doc": path.name, "doc_part": "heading"},
        )
        edges += 1
        codes = _commodities_for_heading(conn, heading4)
        links += link_edge_to_commodities(conn, edge_id, codes)
    return edges, links


def ingest_sections(conn, path: Path) -> tuple[int, int]:
    sections = parse_sections_doc(path)
    edges = 0
    links = 0
    for roman, (title, body) in sections.items():
        if not body.strip():
            continue
        edge_id = f"hsen:section_{roman.lower()}:general"
        upsert_edge(
            conn, edge_id, "hsen_section_general", f"section:{roman}",
            f"HSEN Section {roman} - {title[:100]}",
            body,
            "hsen:sections",
            HSEN_TIER,
            {"source_type": "hsen", "section": roman, "doc": path.name,
             "doc_part": "section_general"},
        )
        edges += 1
        codes = _commodities_for_section(conn, roman)
        links += link_edge_to_commodities(conn, edge_id, codes)
    return edges, links


def main():
    if not HSEN_DIR.exists():
        print(f"HSEN dir not found: {HSEN_DIR}", file=sys.stderr)
        sys.exit(1)
    only_chapters = set(os.environ.get("HSEN_CHAPTERS", "").split(",")) if os.environ.get("HSEN_CHAPTERS") else None
    if only_chapters and "" in only_chapters:
        only_chapters.discard("")
    print(f"only_chapters: {only_chapters or 'ALL'}")

    conn = psycopg.connect(DSN, row_factory=dict_row)
    with_actor(conn, "seeder:hsen")
    try:
        total_edges = 0
        total_links = 0
        files = sorted(HSEN_DIR.glob("Processed downloaded explanatory notes - *.docx"))
        for path in files:
            name = path.stem
            if "sections" in name.lower():
                print(f"\n[sections] {path.name}")
                e, l = ingest_sections(conn, path)
                conn.commit()
                print(f"  {e} edges, {l} links")
                total_edges += e
                total_links += l
                continue
            m = re.search(r"chapter\s+(\d{1,2})", name, re.IGNORECASE)
            if not m:
                continue
            ch2 = m.group(1).zfill(2)
            if only_chapters and ch2 not in only_chapters:
                continue
            print(f"\n[ch{ch2}] {path.name}")
            e, l = ingest_chapter(conn, path)
            conn.commit()
            print(f"  {e} edges, {l} links")
            total_edges += e
            total_links += l
        print(f"\nTotal: {total_edges} edges, {total_links} links inserted")
    finally:
        conn.close()


if __name__ == "__main__":
    main()
